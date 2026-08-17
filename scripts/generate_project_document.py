from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "DOCUMENT_HuongDan_KTX_TayDo.docx"
GITHUB_URL = "https://github.com/HuuNam3/ktx-management-system"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_paragraph(document, text="", bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        prefix.font.name = "Times New Roman"
        prefix.font.size = Pt(12)
        rest = paragraph.add_run(text[len(bold_prefix):])
        rest.font.name = "Times New Roman"
        rest.font.size = Pt(12)
    else:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_info_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    set_cell_text(header[0], "Hạng mục", True)
    set_cell_text(header[1], "Thông tin", True)
    set_cell_shading(header[0], "D9EAF7")
    set_cell_shading(header[1], "D9EAF7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value)
    document.add_paragraph()


def add_code_block(document, lines):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F2F2F2")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    document.add_paragraph()


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENT HƯỚNG DẪN CÀI ĐẶT VÀ SỬ DỤNG")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HỆ THỐNG QUẢN LÝ KÝ TÚC XÁ KTX TÂY ĐÔ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)

    document.add_paragraph()
    add_info_table(
        document,
        [
            ("Tên sản phẩm", "Hệ thống quản lý ký túc xá KTX Tây Đô"),
            ("Phạm vi", "Quản lý lễ tân, phòng, khách/sinh viên, thanh toán, điện nước, hồ sơ 238, giấy tờ, checklist, sự cố, nhà thầu và phản hồi"),
            ("Thiết kế giao diện", "Giao diện được xây dựng dựa trên bản thiết kế Figma được giao"),
            ("Công nghệ frontend", "Angular, Ionic, Tailwind CSS"),
            ("Công nghệ backend", "ASP.NET Core Web API, Entity Framework Core"),
            ("Cơ sở dữ liệu", "PostgreSQL chạy bằng Docker"),
            ("GitHub source", GITHUB_URL),
        ],
    )

    add_heading(document, "1. Giới thiệu hệ thống", 1)
    add_paragraph(
        document,
        "Hệ thống quản lý ký túc xá KTX Tây Đô là sản phẩm hỗ trợ nhân viên quản lý các nghiệp vụ vận hành ký túc xá. "
        "Hệ thống tập trung vào việc số hóa quy trình tiếp nhận, theo dõi khách lưu trú, quản lý phòng, xử lý công nợ, "
        "ghi nhận điện nước, in giấy tờ, quản lý checklist ca trực, sự cố, nhà thầu và phản hồi từ sinh viên."
    )
    add_paragraph(
        document,
        "Sản phẩm được phát triển theo yêu cầu được giao, có sử dụng bản thiết kế Figma làm cơ sở để xây dựng giao diện. "
        "Các chức năng được tổ chức theo từng màn hình quản trị để thuận tiện cho quá trình sử dụng và kiểm thử."
    )

    add_heading(document, "2. Môi trường yêu cầu", 1)
    add_bullet(document, "Hệ điều hành: Windows 10 hoặc Windows 11.")
    add_bullet(document, "Docker Desktop để chạy PostgreSQL.")
    add_bullet(document, ".NET SDK 9 để chạy backend ASP.NET Core.")
    add_bullet(document, "Node.js và npm để chạy frontend Angular.")
    add_bullet(document, "Trình duyệt Chrome, Edge hoặc trình duyệt tương đương.")

    add_heading(document, "3. Cấu trúc source", 1)
    add_paragraph(document, "Mã nguồn sản phẩm được quản lý trên GitHub tại địa chỉ:")
    add_bullet(document, GITHUB_URL)
    add_paragraph(document, "Có thể tải source bằng lệnh:")
    add_code_block(
        document,
        [
            "git clone https://github.com/HuuNam3/ktx-management-system.git",
            "cd ktx-management-system",
        ],
    )
    add_code_block(
        document,
        [
            "ktx-management-system/",
            "├── be/                              # Backend ASP.NET Core Web API",
            "│   └── src/DormitoryManagement.Api/",
            "├── fe/                              # Frontend Angular/Ionic",
            "├── scripts/                         # Script tạo tài liệu/báo cáo",
            "├── BaoCao_ThucTap_KTX_TayDo.docx",
            "└── DOCUMENT_HuongDan_KTX_TayDo.docx",
        ],
    )

    add_heading(document, "4. Cài đặt database bằng Docker Desktop", 1)
    add_paragraph(document, "Database sử dụng PostgreSQL. Có thể tạo container mới trong Docker Desktop theo các bước sau:")
    for step in [
        "Mở Docker Desktop.",
        "Vào mục Images và chọn image postgres:16.",
        "Bấm Run để tạo container mới.",
        "Nhập Container name là next_postgres.",
        "Cấu hình port: Host port 5432 và Container port 5432.",
        "Thêm các biến môi trường POSTGRES_USER, POSTGRES_PASSWORD và POSTGRES_DB.",
        "Bấm Run và kiểm tra container ở trạng thái Running.",
    ]:
        add_number(document, step)
    add_info_table(
        document,
        [
            ("POSTGRES_USER", "admin"),
            ("POSTGRES_PASSWORD", "123456"),
            ("POSTGRES_DB", "nextdb"),
            ("Port", "5432:5432"),
        ],
    )
    add_paragraph(
        document,
        "Nếu database cũ bị sai cấu trúc bảng, cần xóa container PostgreSQL đang dùng trước, sau đó xóa volume dữ liệu cũ trong mục Volumes của Docker Desktop rồi tạo lại container."
    )

    add_heading(document, "5. Chạy backend", 1)
    add_paragraph(document, "Backend nằm trong thư mục be/src/DormitoryManagement.Api. Chạy backend bằng PowerShell hoặc terminal trong VS Code:")
    add_code_block(
        document,
        [
            "cd D:\\project\\Angular\\ktx-management-system\\be\\src\\DormitoryManagement.Api",
            "dotnet run",
        ],
    )
    add_paragraph(document, "Sau khi chạy thành công, backend sử dụng địa chỉ:")
    add_bullet(document, "API: http://localhost:5177")
    add_bullet(document, "Swagger: http://localhost:5177/swagger/index.html")
    add_paragraph(
        document,
        "Khi backend khởi động, hệ thống tự tạo bảng và dữ liệu mẫu bằng Entity Framework Core nếu database đang trống."
    )

    add_heading(document, "6. Chạy frontend", 1)
    add_paragraph(document, "Frontend nằm trong thư mục fe. Chạy frontend bằng lệnh:")
    add_code_block(
        document,
        [
            "cd D:\\project\\Angular\\ktx-management-system\\fe",
            "npm install",
            "npm run start -- --host 127.0.0.1 --port 4200",
        ],
    )
    add_paragraph(document, "Sau khi chạy thành công, mở trình duyệt tại địa chỉ:")
    add_bullet(document, "Frontend: http://127.0.0.1:4200/")

    add_heading(document, "7. Tài khoản đăng nhập", 1)
    add_info_table(
        document,
        [
            ("Tên đăng nhập", "admin"),
            ("Mật khẩu", "admin"),
            ("Vai trò", "Quản trị viên"),
        ],
    )

    add_heading(document, "8. Kiểm tra kết nối backend và database", 1)
    add_paragraph(document, "Sau khi chạy Docker và backend, có thể kiểm tra bằng Swagger hoặc mở trực tiếp endpoint thống kê:")
    add_bullet(document, "http://localhost:5177/api/reception-logs/stats")
    add_paragraph(document, "Nếu API trả về dữ liệu JSON, backend đã kết nối thành công với PostgreSQL.")
    add_code_block(
        document,
        [
            "{",
            '  "todayCheckins": 45,',
            '  "todayCheckouts": 23,',
            '  "totalEntries": 68,',
            '  "totalExits": 32',
            "}",
        ],
    )

    add_heading(document, "9. Danh sách chức năng chính", 1)
    features = [
        ("Đăng nhập", "Cho phép người dùng đăng nhập vào hệ thống quản trị bằng tài khoản được cấp."),
        ("Tổng quan", "Hiển thị số liệu tổng hợp về khách đang ở, phòng trống, doanh thu và cảnh báo."),
        ("Lễ tân", "Tìm kiếm khách, check-in, check-out và quét QR cho khách/sinh viên."),
        ("Quản lý phòng", "Hiển thị sơ đồ phòng, số giường, trạng thái còn trống, còn chỗ, đầy hoặc bảo trì."),
        ("Nhật ký ra vào", "Theo dõi lịch sử check-in/check-out, lọc theo loại giao dịch và ngày."),
        ("Quản lý khách", "Quản lý danh sách khách/sinh viên, thông tin CCCD, số điện thoại, phòng, công nợ và trạng thái."),
        ("Quản lý vé", "Quản lý vé ăn, vé xe, doanh thu và giao dịch vé gần đây."),
        ("Thanh toán", "Theo dõi công nợ, trạng thái đã thanh toán, quá hạn và xử lý thu tiền."),
        ("Điện nước", "Nhập chỉ số điện nước, tính toán tiêu thụ và trạng thái thu tiền."),
        ("Hồ sơ 238", "Duyệt hồ sơ miễn giảm, xác nhận, từ chối và theo dõi trạng thái hồ sơ."),
        ("In giấy tờ", "Tạo các loại giấy xác nhận, hóa đơn, biên lai và hợp đồng lưu trú."),
        ("Checklist", "Theo dõi công việc ca trực, đánh dấu hoàn thành và ghi chú ca trực."),
        ("Sự cố", "Ghi nhận sự cố, chi phí bồi thường và xử lý trạng thái."),
        ("Nhà thầu/Thợ", "Check-in/check-out thợ, theo dõi công việc và trạng thái làm việc."),
        ("Phản hồi", "Quản lý phản hồi, khiếu nại, mức độ ưu tiên và trạng thái xử lý."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], "Chức năng", True)
    set_cell_text(table.rows[0].cells[1], "Mô tả", True)
    set_cell_shading(table.rows[0].cells[0], "D9EAF7")
    set_cell_shading(table.rows[0].cells[1], "D9EAF7")
    for name, description in features:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, True)
        set_cell_text(cells[1], description)
    document.add_paragraph()

    add_heading(document, "10. Một số API sử dụng trong hệ thống", 1)
    add_info_table(
        document,
        [
            ("POST /api/auth/login", "Đăng nhập và nhận token truy cập."),
            ("POST /api/auth/refresh", "Làm mới token đăng nhập."),
            ("GET /api/reception-logs/stats", "Lấy thống kê check-in/check-out."),
            ("GET /api/reception-logs", "Lấy danh sách nhật ký ra vào."),
            ("GET /api/reports/dashboard", "Lấy dữ liệu tổng quan hệ thống."),
        ],
    )

    add_heading(document, "11. Quy trình sử dụng cơ bản", 1)
    for step in [
        "Khởi động Docker Desktop và kiểm tra container PostgreSQL đang chạy.",
        "Chạy backend ASP.NET Core tại cổng 5177.",
        "Chạy frontend Angular tại cổng 4200.",
        "Đăng nhập bằng tài khoản admin.",
        "Vào từng menu quản trị để thực hiện nghiệp vụ tương ứng.",
        "Kiểm tra dữ liệu thay đổi trên màn hình và qua API Swagger khi cần.",
    ]:
        add_number(document, step)

    add_heading(document, "12. Ghi chú khi gặp lỗi", 1)
    add_bullet(document, "Nếu Swagger báo lỗi relation \"user\" does not exist, cần xóa container và volume PostgreSQL cũ rồi tạo lại database sạch.")
    add_bullet(document, "Nếu port 5432 bị chiếm, kiểm tra container PostgreSQL khác đang chạy và dừng container không sử dụng.")
    add_bullet(document, "Nếu frontend vẫn hiển thị giao diện cũ, nhấn Ctrl + F5 để tải lại bundle mới.")
    add_bullet(document, "Nếu backend không chạy được do port 5177 bị chiếm, dừng process backend cũ rồi chạy lại.")

    add_heading(document, "13. Kết luận", 1)
    add_paragraph(
        document,
        "Tài liệu này cung cấp thông tin cần thiết để cài đặt, khởi chạy và sử dụng hệ thống quản lý ký túc xá KTX Tây Đô. "
        "Sản phẩm bao gồm đầy đủ các màn hình quản trị chính, dữ liệu mẫu phục vụ kiểm thử và hướng dẫn vận hành bằng Docker, backend và frontend."
    )

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "DOCUMENT - Hệ thống quản lý ký túc xá KTX Tây Đô"

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
