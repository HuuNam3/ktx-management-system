from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "BaoCao_ThucTap_KTX_TayDo.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(13)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 13, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold


def p(doc, text="", bold=False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    return paragraph


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def page_break(doc):
    doc.add_page_break()


def cover(doc):
    p(doc, "ỦY BAN NHÂN DÂN TP HỒ CHÍ MINH", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "TRƯỜNG ĐẠI HỌC SÀI GÒN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")
    p(doc, "")
    p(doc, "Họ và tên sinh viên: [...]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "MSSV: [...]        Lớp: [...]", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")
    p(doc, "")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BÁO CÁO\nTHỰC TẬP TỐT NGHIỆP")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(22)
    p(doc, "")
    p(doc, "Đề tài: XÂY DỰNG HỆ THỐNG QUẢN LÝ KÝ TÚC XÁ KTX TÂY ĐÔ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")
    p(doc, "Công ty thực tập        : [...]", align=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "Chuyên gia hướng dẫn    : [...]", align=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "Giảng viên hướng dẫn    : [...]", align=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "")
    p(doc, "")
    p(doc, "TP. Hồ Chí Minh, tháng [...] năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def remarks(doc, title):
    heading(doc, title, 1)
    for _ in range(13):
        p(doc, "." * 120)
    p(doc, "")
    p(doc, "TP. Hồ Chí Minh, ngày ..... tháng ..... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, "Người nhận xét", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, "(Ký và ghi rõ họ tên)", align=WD_ALIGN_PARAGRAPH.RIGHT)
    page_break(doc)


def toc(doc):
    heading(doc, "MỤC LỤC", 1)
    items = [
        "LỜI MỞ ĐẦU",
        "CHƯƠNG 1. GIỚI THIỆU",
        "1.1. Giới thiệu đơn vị thực tập",
        "1.2. Nhiệm vụ thực tập",
        "1.3. Kết luận chương 1",
        "CHƯƠNG 2. PHÂN TÍCH VÀ XÂY DỰNG HỆ THỐNG QUẢN LÝ KÝ TÚC XÁ",
        "2.1. Tổng quan bài toán",
        "2.2. Công nghệ sử dụng",
        "2.3. Phân tích chức năng hệ thống",
        "2.4. Thiết kế cơ sở dữ liệu",
        "2.5. Thiết kế giao diện và trải nghiệm người dùng",
        "2.6. Cài đặt frontend",
        "2.7. Cài đặt backend và API",
        "2.8. Kiểm thử và kết quả thực nghiệm",
        "CHƯƠNG 3. KẾT QUẢ THỰC TẬP",
        "3.1. Kết quả đạt được",
        "3.2. Bảng ghi nhận kết quả thực tập hằng tuần",
        "3.3. Đánh giá cá nhân",
        "CHƯƠNG 4. KẾT LUẬN VÀ KIẾN NGHỊ",
        "TÀI LIỆU THAM KHẢO",
        "PHỤ LỤC",
    ]
    for item in items:
        p(doc, item)
    p(doc, "Ghi chú: Sau khi mở bằng Microsoft Word, có thể thay phần này bằng mục lục tự động bằng chức năng References -> Table of Contents.")
    page_break(doc)


def intro(doc):
    heading(doc, "LỜI MỞ ĐẦU", 1)
    paragraphs = [
        "Trong bối cảnh chuyển đổi số đang diễn ra mạnh mẽ, việc ứng dụng công nghệ thông tin vào công tác quản lý nội bộ trở thành nhu cầu cần thiết đối với nhiều đơn vị giáo dục, doanh nghiệp và tổ chức dịch vụ. Đối với ký túc xá, khối lượng công việc hằng ngày khá lớn, bao gồm quản lý thông tin sinh viên, quản lý phòng ở, theo dõi check-in/check-out, ghi nhận điện nước, thanh toán công nợ, xử lý sự cố, phản hồi và các loại hồ sơ liên quan. Nếu các nghiệp vụ này được thực hiện thủ công bằng giấy tờ hoặc bảng tính rời rạc thì dễ xảy ra sai sót, khó thống kê và mất nhiều thời gian tra cứu.",
        "Trong thời gian thực tập tốt nghiệp, em được giao thực hiện đề tài “Xây dựng hệ thống quản lý ký túc xá KTX Tây Đô”. Đề tài có sẵn định hướng giao diện trên Figma, vì vậy nhiệm vụ của em là phân tích thiết kế Figma, chuyển đổi thành giao diện web hoạt động được, bổ sung các tương tác cần thiết và xây dựng phần backend/API phục vụ lưu trữ, xử lý dữ liệu. Hệ thống được xây dựng theo mô hình web application, trong đó frontend sử dụng Angular, Ionic và Tailwind CSS; backend sử dụng ASP.NET Core Web API, Entity Framework Core và PostgreSQL. Sản phẩm hướng đến mục tiêu hỗ trợ nhân viên ký túc xá thao tác nhanh, giảm nhầm lẫn trong quá trình vận hành và tạo nền tảng để mở rộng thành hệ thống quản lý hoàn chỉnh.",
        "Báo cáo này trình bày quá trình thực hiện trong 8 tuần, bao gồm tìm hiểu yêu cầu, phân tích nghiệp vụ, thiết kế giao diện, xây dựng chức năng, cài đặt backend, tạo dữ liệu mẫu và kiểm thử hệ thống. Bên cạnh những kết quả đã đạt được, báo cáo cũng nêu ra các hạn chế còn tồn tại và đề xuất hướng phát triển trong tương lai.",
    ]
    for text in paragraphs:
        p(doc, text)
    page_break(doc)


def chapter1(doc):
    heading(doc, "CHƯƠNG 1. GIỚI THIỆU", 1)
    heading(doc, "1.1. Giới thiệu đơn vị thực tập", 2)
    p(doc, "Đơn vị thực tập: [...].")
    p(doc, "Địa chỉ: [...].")
    p(doc, "Điện thoại/Email: [...].")
    p(doc, "Lĩnh vực hoạt động: [...].")
    p(doc, "Trong thời gian thực tập 8 tuần, sinh viên được tiếp cận môi trường làm việc liên quan đến phát triển phần mềm, tìm hiểu quy trình phân tích yêu cầu, thiết kế giao diện, xây dựng ứng dụng web, kiểm thử chức năng và hoàn thiện sản phẩm. Nội dung thực tập tập trung vào việc vận dụng kiến thức đã học vào một bài toán quản lý thực tế.")
    heading(doc, "1.2. Nhiệm vụ thực tập", 2)
    p(doc, "Các nhiệm vụ chính được thực hiện trong quá trình thực tập 8 tuần gồm:")
    for item in [
        "Khảo sát bài toán quản lý ký túc xá, xác định các nhóm chức năng chính và dữ liệu cần quản lý.",
        "Phân tích file thiết kế Figma được cung cấp, chuyển đổi các màn hình thiết kế thành giao diện web có thể thao tác.",
        "Cài đặt frontend bằng Angular/Ionic, tổ chức routing, guard đăng nhập, component dùng lại và các màn hình nghiệp vụ.",
        "Cài đặt tương tác frontend cho sản phẩm như đăng nhập, tìm kiếm, lọc dữ liệu, xem/sửa thông tin, lưu ghi chú và hiển thị thông báo thao tác.",
        "Xây dựng backend ASP.NET Core Web API, cấu hình CORS, JWT, Entity Framework Core và PostgreSQL.",
        "Tạo dữ liệu mẫu trong database seeder để phục vụ quá trình báo cáo và trình diễn hệ thống.",
        "Kiểm thử build frontend/backend, ghi nhận lỗi phát sinh và điều chỉnh giao diện cho phù hợp kích thước hiển thị.",
    ]:
        bullet(doc, item)
    heading(doc, "1.3. Kết luận chương 1", 2)
    p(doc, "Chương 1 đã trình bày khái quát về bối cảnh thực tập, đơn vị thực tập và các nhiệm vụ được giao. Những nội dung này là cơ sở để triển khai phân tích, thiết kế và cài đặt hệ thống quản lý ký túc xá trong các chương tiếp theo.")
    page_break(doc)


def chapter2(doc):
    heading(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ XÂY DỰNG HỆ THỐNG QUẢN LÝ KÝ TÚC XÁ", 1)
    heading(doc, "2.1. Tổng quan bài toán", 2)
    p(doc, "Hệ thống quản lý ký túc xá KTX Tây Đô được xây dựng nhằm hỗ trợ nhân viên quản lý các nghiệp vụ vận hành thường ngày. Người dùng chính của hệ thống là nhân viên lễ tân, nhân viên quản lý phòng, bộ phận thanh toán và người quản trị. Các nghiệp vụ được gom thành các phân hệ rõ ràng để người dùng có thể tra cứu, xử lý và theo dõi thông tin nhanh chóng.")
    p(doc, "Các chức năng chính của hệ thống gồm: đăng nhập, tổng quan thống kê, quản lý lễ tân, quản lý phòng, nhật ký ra vào, quản lý khách và sinh viên, quản lý vé, thanh toán, điện nước, hồ sơ 238, in giấy tờ, checklist nhân viên, sự cố và bồi thường, nhà thầu/thợ, phản hồi và khiếu nại.")
    heading(doc, "2.2. Công nghệ sử dụng", 2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Thành phần", True)
    set_cell_text(hdr[1], "Công nghệ", True)
    set_cell_text(hdr[2], "Vai trò", True)
    rows = [
        ("Frontend", "Angular 20, Ionic 8, TypeScript, RxJS", "Xây dựng giao diện, routing, component và tương tác người dùng."),
        ("Giao diện", "Tailwind CSS, Ionicons, Lucide Icons", "Thiết kế dashboard, bảng dữ liệu, thẻ thống kê, icon và trạng thái."),
        ("Backend", "ASP.NET Core Web API .NET 9", "Xây dựng API, xử lý nghiệp vụ và cung cấp dữ liệu cho frontend."),
        ("Database", "PostgreSQL, Entity Framework Core, Npgsql", "Lưu trữ dữ liệu phòng, khách, hợp đồng, hóa đơn, điện nước, sự cố."),
        ("Bảo mật", "JWT Bearer Authentication, BCrypt", "Đăng nhập, sinh token và mã hóa mật khẩu."),
        ("Công cụ", "Swagger, Angular CLI, .NET CLI", "Kiểm thử API, build và chạy ứng dụng."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    heading(doc, "2.3. Phân tích chức năng hệ thống", 2)
    heading(doc, "2.3.1. Chức năng đăng nhập", 3)
    p(doc, "Màn hình đăng nhập cho phép người dùng nhập tài khoản và mật khẩu để truy cập hệ thống. Frontend có cơ chế Auth Guard nhằm ngăn người dùng chưa đăng nhập truy cập trực tiếp vào các trang quản trị. Hệ thống hỗ trợ tài khoản đăng nhập và cơ chế kiểm soát truy cập; backend cũng có cấu trúc AuthController, JWT và refresh token để phát triển đăng nhập thật.")
    heading(doc, "2.3.2. Chức năng tổng quan", 3)
    p(doc, "Màn hình tổng quan cung cấp các chỉ số nhanh như tổng số khách, số phòng trống, doanh thu hôm nay và cảnh báo. Ngoài ra, màn hình còn có khu vực biểu đồ doanh thu 7 ngày, phân bổ loại phòng, hoạt động gần đây và danh sách cảnh báo/thông báo.")
    heading(doc, "2.3.3. Chức năng lễ tân và nhật ký ra vào", 3)
    p(doc, "Phân hệ lễ tân hỗ trợ tìm kiếm nhanh khách theo CCCD, tên hoặc số phòng, thực hiện check-in, check-out và quét QR. Phân hệ nhật ký ra vào hiển thị lịch sử giao dịch theo thời gian, loại giao dịch, họ tên, CCCD, phòng, phương thức và nhân viên xử lý. Người dùng có thể tìm kiếm, lọc theo loại giao dịch và xuất báo cáo.")
    heading(doc, "2.3.4. Chức năng quản lý phòng", 3)
    p(doc, "Màn hình quản lý phòng thể hiện sơ đồ phòng và trạng thái giường. Các trạng thái chính gồm còn trống, còn chỗ, đầy và bảo trì. Giao diện có thẻ thống kê tổng số phòng, phòng còn trống, phòng còn chỗ, phòng đầy và phòng bảo trì, đồng thời hỗ trợ lọc theo dãy A, B, C.")
    heading(doc, "2.3.5. Chức năng quản lý khách và sinh viên", 3)
    p(doc, "Phân hệ quản lý khách và sinh viên cho phép theo dõi danh sách người đang ở ký túc xá. Bảng dữ liệu gồm họ tên, CCCD, số điện thoại, phòng - giường, loại đối tượng, ngày vào, công nợ và trạng thái. Giao diện hỗ trợ tìm kiếm và xem/sửa thông tin ở mức giao diện sản phẩm.")
    heading(doc, "2.3.6. Chức năng vé, thanh toán và điện nước", 3)
    p(doc, "Phân hệ quản lý vé theo dõi vé ăn, vé xe, doanh thu tháng và suất ăn trong ngày. Phân hệ thanh toán quản lý công nợ, trạng thái quá hạn, chưa đóng, đã đóng một phần và thao tác thanh toán. Phân hệ điện nước hỗ trợ nhập chỉ số, tính tiêu thụ điện/nước, tổng tiền và trạng thái đã thu hoặc chưa thu.")
    heading(doc, "2.3.7. Các chức năng hỗ trợ vận hành", 3)
    p(doc, "Các chức năng hỗ trợ gồm quản lý hồ sơ 238, in giấy tờ, checklist nhân viên, sự cố và bồi thường, nhà thầu/thợ, phản hồi và khiếu nại. Đây là các nghiệp vụ giúp ký túc xá theo dõi công việc nội bộ, xử lý yêu cầu và chuẩn bị tài liệu cần thiết cho sinh viên.")

    heading(doc, "2.4. Thiết kế cơ sở dữ liệu", 2)
    p(doc, "Backend sử dụng Entity Framework Core để ánh xạ các entity sang bảng trong PostgreSQL. Các nhóm bảng chính gồm:")
    for item in [
        "Nhóm người dùng và phân quyền: Users, Roles, RefreshTokens, AuditLogs.",
        "Nhóm cơ sở vật chất: Buildings, RoomTypes, Rooms, Beds, Assets, RoomAssets.",
        "Nhóm khách và hợp đồng: Tenants, Contracts, ContractTenants.",
        "Nhóm dịch vụ và điện nước: UtilityServices, Meters, MeterReadings.",
        "Nhóm hóa đơn và thanh toán: Invoices, InvoiceItems, Payments.",
        "Nhóm vận hành: MaintenanceRequests, Notifications, NotificationRecipients.",
    ]:
        bullet(doc, item)
    p(doc, "Dữ liệu mẫu đã được tạo trong DatabaseSeeder với các phòng A201, A202, A203, A204, A205, A301, B205, B305, C108; các khách/sinh viên như Nguyễn Văn A, Trần Thị B, Lê Văn C, Phạm Thị D, Hoàng Văn E; đồng thời có dữ liệu công nợ, điện nước, sự cố, thông báo và nhật ký hoạt động.")

    heading(doc, "2.5. Thiết kế giao diện và trải nghiệm người dùng", 2)
    p(doc, "Giao diện hệ thống được triển khai dựa trên mẫu thiết kế Figma được giao. Thiết kế có dạng admin dashboard với thanh điều hướng bên trái, thanh thông tin phía trên và vùng nội dung chính. Màu sắc chủ đạo là trắng, xám nhạt và xanh dương, kết hợp các màu trạng thái như xanh lá cho thành công, đỏ cho cảnh báo, vàng cho chờ xử lý và tím cho chức năng QR hoặc hồ sơ đặc biệt. Trong quá trình cài đặt, em dựa vào Figma và các ảnh chụp màn hình từ thiết kế để dựng lại layout, khoảng cách, kích thước chữ, thẻ thống kê, bảng dữ liệu và trạng thái hiển thị.")
    p(doc, "Trong quá trình hoàn thiện, giao diện đã được chỉnh để tránh lỗi tràn ngang ở màn Nhật ký ra vào, đặc biệt tại khu vực thẻ thống kê và bảng dữ liệu. Wrapper chính của ứng dụng được điều chỉnh để chiều rộng nội dung tính đúng theo sidebar, giúp các card không bị cắt ở mép phải.")

    heading(doc, "2.6. Cài đặt frontend", 2)
    p(doc, "Frontend được tổ chức theo kiến trúc component và route. Mỗi màn hình nghiệp vụ nằm trong thư mục pages, các thành phần dùng chung như sidebar, topbar, page-header, search-card, statistic-card và room-card nằm trong thư mục components. Các service chịu trách nhiệm cung cấp dữ liệu mẫu hoặc gọi API backend.")
    p(doc, "Danh sách route chính gồm: /dang-nhap, /tong-quan, /reception, /phong, /nhat-ky, /khach, /ve, /thanh-toan, /dien-nuoc, /ho-so, /in-giay-to, /checklist, /su-co, /nha-thau và /phan-hoi.")
    p(doc, "Các tương tác frontend đã cài đặt gồm tìm kiếm, lọc dữ liệu, mở panel chi tiết, sửa thông tin, lưu ghi chú, xử lý sự cố/phản hồi, check-out thợ, xuất báo cáo và thông báo toast.")

    heading(doc, "2.7. Cài đặt backend và API", 2)
    p(doc, "Backend sử dụng ASP.NET Core Web API. Program.cs cấu hình controller, Swagger, JWT Bearer, CORS, DbContext PostgreSQL, repository/service dùng chung và middleware xử lý ngoại lệ. Các controller CRUD chung được xây dựng cho nhiều entity như Rooms, Tenants, Invoices, Payments, MaintenanceRequests, Notifications.")
    p(doc, "Ngoài các controller CRUD, hệ thống có AuthController cho đăng nhập và ReceptionLogsController phục vụ phân hệ nhật ký ra vào. API nhật ký gồm:")
    for item in [
        "GET /api/reception-logs/stats: lấy số liệu check-in, check-out, tổng giao dịch và số lượt QR.",
        "GET /api/reception-logs?type=checkin&search=A301: lấy danh sách nhật ký theo loại và từ khóa tìm kiếm.",
        "POST /api/reception-logs: tạo một bản ghi nhật ký mới ở mức dữ liệu mẫu.",
    ]:
        bullet(doc, item)
    p(doc, "Frontend đã được cấu hình HttpClient và environment.apiUrl để gọi API backend tại http://localhost:5177/api. Riêng LogService có cơ chế fallback dữ liệu mẫu nếu backend chưa được bật, giúp quá trình trình diễn không bị gián đoạn.")

    heading(doc, "2.8. Kiểm thử và kết quả thực nghiệm", 2)
    p(doc, "Quá trình kiểm thử tập trung vào việc kiểm tra build, kiểm tra routing, kiểm tra giao diện ở kích thước trình duyệt và kiểm tra API backend. Các lệnh đã sử dụng gồm:")
    for item in [
        "npx tsc -p tsconfig.app.json --noEmit: kiểm tra TypeScript và template Angular.",
        "npm run build: build production frontend.",
        "dotnet build DormitoryManagement.sln --no-restore: kiểm tra backend .NET.",
        "Gọi thử API /api/reception-logs/stats và /api/reception-logs?type=checkin&search=A301 để xác nhận backend trả dữ liệu.",
    ]:
        bullet(doc, item)
    p(doc, "Kết quả kiểm thử cho thấy frontend và backend đều build thành công. Backend có thể chạy API ngay cả khi PostgreSQL chưa được bật; khi PostgreSQL được cấu hình đúng, DatabaseSeeder sẽ tạo dữ liệu mẫu cho database.")
    heading(doc, "2.9. Kết luận chương 2", 2)
    p(doc, "Chương 2 đã trình bày quá trình phân tích và xây dựng hệ thống quản lý ký túc xá từ giao diện, chức năng, cơ sở dữ liệu đến backend API. Hệ thống đã có nền tảng tương đối đầy đủ để vận hành và có thể tiếp tục mở rộng thành sản phẩm hoàn chỉnh.")
    page_break(doc)


def chapter3(doc):
    heading(doc, "CHƯƠNG 3. KẾT QUẢ THỰC TẬP", 1)
    heading(doc, "3.1. Kết quả đạt được", 2)
    for item in [
        "Hoàn thiện giao diện admin dashboard cho hệ thống quản lý KTX Tây Đô dựa trên thiết kế Figma được giao.",
        "Cài đặt các màn hình nghiệp vụ quan trọng: tổng quan, lễ tân, quản lý phòng, nhật ký ra vào, quản lý khách, vé, thanh toán, điện nước, hồ sơ 238, in giấy tờ, checklist, sự cố, nhà thầu và phản hồi.",
        "Cài đặt đăng nhập và guard bảo vệ route.",
        "Cài đặt tương tác frontend phục vụ trình diễn: tìm kiếm, lọc, xem/sửa, lưu ghi chú, xử lý và thông báo toast.",
        "Xây dựng backend ASP.NET Core Web API với các entity, DbContext, repository/service, controller CRUD và controller nhật ký ra vào.",
        "Tạo dữ liệu mẫu trong seeder cho các bảng chính để phục vụ dữ liệu database.",
        "Kiểm tra build frontend và backend thành công.",
    ]:
        bullet(doc, item)
    heading(doc, "3.2. Bảng ghi nhận kết quả thực tập hằng tuần", 2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    for i, title in enumerate(["Tuần", "Nội dung công việc", "Kết quả đạt được", "Ghi chú"]):
        set_cell_text(table.rows[0].cells[i], title, True)
    weekly = [
        ("Tuần 1", "Khảo sát yêu cầu, tìm hiểu nghiệp vụ quản lý ký túc xá, xác định các phân hệ chính.", "Xây dựng danh sách chức năng và dữ liệu cần quản lý.", ""),
        ("Tuần 2", "Thiết kế cấu trúc frontend, routing, layout sidebar/topbar và trang đăng nhập.", "Có khung ứng dụng, điều hướng và đăng nhập.", ""),
        ("Tuần 3", "Xây dựng các màn tổng quan, lễ tân, quản lý phòng và nhật ký ra vào.", "Hoàn thiện các màn cốt lõi phục vụ vận hành.", ""),
        ("Tuần 4", "Xây dựng các màn quản lý khách, vé, thanh toán, điện nước và hồ sơ 238.", "Có giao diện dữ liệu dạng bảng, thẻ thống kê và bộ lọc.", ""),
        ("Tuần 5", "Xây dựng các màn in giấy tờ, checklist, sự cố, nhà thầu và phản hồi.", "Hoàn thiện các phân hệ hỗ trợ vận hành.", ""),
        ("Tuần 6", "Cài đặt backend .NET API, DbContext, JWT, controller và seeder dữ liệu mẫu.", "Backend build thành công, có API cho nhật ký ra vào.", ""),
        ("Tuần 7", "Kiểm thử giao diện, sửa lỗi tràn layout, kiểm tra build frontend/backend.", "Ứng dụng chạy ổn định ở môi trường local.", ""),
        ("Tuần 8", "Tổng hợp tài liệu, viết báo cáo và chuẩn bị nội dung trình bày.", "Hoàn thiện bản báo cáo thực tập và nội dung trình bày sản phẩm.", ""),
    ]
    for row in weekly:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    heading(doc, "3.3. Đánh giá cá nhân", 2)
    p(doc, "Qua quá trình thực hiện đề tài được giao, em đã củng cố kỹ năng đọc và phân tích thiết kế Figma, chuyển đổi giao diện thiết kế thành mã nguồn Angular, tổ chức component/service, xây dựng API bằng ASP.NET Core và thiết kế dữ liệu bằng Entity Framework Core. Bên cạnh đó, em cũng rèn luyện khả năng xử lý lỗi phát sinh trong quá trình build, kiểm thử giao diện, điều chỉnh responsive layout và chuẩn bị dữ liệu mẫu phục vụ báo cáo.")
    heading(doc, "3.4. Kết luận chương 3", 2)
    p(doc, "Chương 3 đã tổng hợp các kết quả đạt được trong quá trình thực tập 8 tuần. Sản phẩm hiện tại đáp ứng tốt các mục tiêu đã được giao và là nền tảng để phát triển thêm các chức năng backend hoàn chỉnh trong tương lai.")
    page_break(doc)


def chapter4(doc):
    heading(doc, "CHƯƠNG 4. KẾT LUẬN VÀ KIẾN NGHỊ", 1)
    heading(doc, "4.1. Kết luận", 2)
    p(doc, "Đề tài “Xây dựng hệ thống quản lý ký túc xá KTX Tây Đô” đã hoàn thành các mục tiêu chính đặt ra ở mức sản phẩm phục vụ báo cáo thực tập. Hệ thống có giao diện quản trị trực quan, nhiều phân hệ nghiệp vụ, tương tác frontend đầy đủ cho quá trình sử dụng và backend API nền tảng. Việc xây dựng dữ liệu mẫu giúp hệ thống có nội dung thực tế hơn, hỗ trợ việc kiểm thử và trình bày kết quả.")
    p(doc, "Thông qua đề tài được giao, sinh viên có cơ hội vận dụng kiến thức đã học vào một bài toán quản lý thực tế, đồng thời tiếp cận quy trình phát triển ứng dụng web hiện đại gồm phân tích yêu cầu, đọc thiết kế Figma, lập trình frontend, xây dựng backend, thiết kế database, tạo dữ liệu mẫu và kiểm thử.")
    heading(doc, "4.2. Hạn chế", 2)
    for item in [
        "Một số chức năng backend mới ở mức dữ liệu mẫu, chưa kết nối đầy đủ tất cả màn hình frontend với database thật.",
        "Chưa hoàn thiện toàn bộ nghiệp vụ phân quyền chi tiết theo từng vai trò.",
        "Chưa có chức năng upload tài liệu, in file PDF thật hoặc xuất Excel đầy đủ.",
        "Chưa triển khai môi trường production và chưa kiểm thử tải lớn.",
    ]:
        bullet(doc, item)
    heading(doc, "4.3. Hướng phát triển", 2)
    for item in [
        "Hoàn thiện CRUD backend cho toàn bộ phân hệ và nối toàn bộ frontend với API thật.",
        "Bổ sung xác thực frontend bằng JWT thật, refresh token và phân quyền theo vai trò.",
        "Thêm chức năng upload hồ sơ, sinh PDF giấy xác nhận, xuất Excel và thống kê biểu đồ từ dữ liệu thật.",
        "Bổ sung kiểm thử tự động, logging, audit trail và dashboard báo cáo nâng cao.",
        "Triển khai hệ thống lên server hoặc cloud để sử dụng thực tế.",
    ]:
        bullet(doc, item)
    heading(doc, "4.4. Kiến nghị", 2)
    p(doc, "Về phía doanh nghiệp/đơn vị thực tập, cần tiếp tục chuẩn hóa quy trình nghiệp vụ và dữ liệu đầu vào để việc số hóa đạt hiệu quả cao. Về phía cơ sở đào tạo, sinh viên mong muốn được tiếp cận nhiều hơn với các tình huống dự án thực tế, đặc biệt là quy trình làm việc nhóm, kiểm thử và triển khai sản phẩm.")
    page_break(doc)


def references_and_appendix(doc):
    heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Microsoft, ASP.NET Core Documentation, https://learn.microsoft.com/aspnet/core.",
        "Microsoft, Entity Framework Core Documentation, https://learn.microsoft.com/ef/core.",
        "Angular Documentation, https://angular.dev.",
        "Ionic Framework Documentation, https://ionicframework.com/docs.",
        "PostgreSQL Documentation, https://www.postgresql.org/docs.",
        "Tailwind CSS Documentation, https://tailwindcss.com/docs.",
    ]
    for ref in refs:
        numbered(doc, ref)
    page_break(doc)

    heading(doc, "PHỤ LỤC", 1)
    heading(doc, "Phụ lục A. Cách chạy hệ thống", 2)
    p(doc, "Chạy frontend:")
    p(doc, "cd fe")
    p(doc, "npm start -- --host 127.0.0.1 --port 4200")
    p(doc, "Địa chỉ truy cập: http://127.0.0.1:4200")
    p(doc, "Tài khoản frontend: admin / admin hoặc admin@ktx.local / 123456")
    p(doc, "Chạy backend:")
    p(doc, "cd be")
    p(doc, "dotnet run --project src/DormitoryManagement.Api/DormitoryManagement.Api.csproj --launch-profile http")
    p(doc, "Địa chỉ API: http://localhost:5177")
    p(doc, "Tài khoản backend seeded: admin / Admin@123")
    heading(doc, "Phụ lục B. Một số API", 2)
    for item in [
        "GET http://localhost:5177/api/reception-logs/stats",
        "GET http://localhost:5177/api/reception-logs?type=checkin&search=A301",
        "POST http://localhost:5177/api/reception-logs",
        "POST http://localhost:5177/api/auth/login",
    ]:
        bullet(doc, item)
    heading(doc, "Phụ lục C. Dữ liệu mẫu", 2)
    p(doc, "DatabaseSeeder tạo dữ liệu mẫu gồm dãy A/B/C, phòng A201, A202, A203, A204, A205, A301, B205, B305, C108; danh sách sinh viên/khách như Nguyễn Văn A, Trần Thị B, Lê Văn C, Phạm Thị D, Hoàng Văn E; dữ liệu công nợ, điện nước, sự cố, thông báo và nhật ký hoạt động.")


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)
    remarks(doc, "NHẬN XÉT CỦA CHUYÊN GIA DOANH NGHIỆP")
    remarks(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    intro(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    references_and_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
